from __future__ import annotations

import csv
import io

import docx
import ftfy
import openpyxl
from pypdf import PdfReader


def parse_csv_rows(content: bytes) -> list[dict]:
    """Parse CSV bytes into row dicts keyed by the file's own header row.

    Decodes with `utf-8-sig` so a file exported from Excel with a leading
    byte-order mark doesn't corrupt the first header name into
    `'\\ufefftitle'`.
    """
    text = content.decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    return [dict(row) for row in reader]


def parse_xlsx_rows(content: bytes) -> list[dict]:
    """Parse XLSX bytes into row dicts keyed by the active sheet's header
    row (its first row). Skips rows where every cell is empty -- a real
    trailing-blank-row guard, not conditioned on any one specific column:
    checking only the first cell (an earlier version of this function did)
    silently drops legitimate data rows whenever the sheet's first column
    holds optional metadata (e.g. "notes"/"date") that happens to be blank
    on some rows while the actual text lives in a later column.
    """
    workbook = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    worksheet = workbook.active
    rows_iter = worksheet.iter_rows(values_only=True)
    try:
        header = next(rows_iter)
    except StopIteration:
        # A blank/empty sheet has no header row at all -- without this,
        # `next()` raises a bare `StopIteration` whose `str()` is empty,
        # producing an uninformative "could not parse file as XLSX: " in
        # the app.py caller's HTTPException detail.
        raise ValueError("worksheet has no header row (the file appears to be empty)") from None
    if any(name is None for name in header):
        # A `None`-named column (an unlabeled header cell) would otherwise
        # produce a `None`-keyed row dict; `app.py`'s `_rows_to_texts` later
        # does `sorted(rows[0].keys())`, and sorting a mix of `str` and
        # `None` raises `TypeError` -- fail with a clear message here
        # instead of that opaque 500 downstream.
        raise ValueError(f"worksheet header row has one or more empty/unnamed cells: {header!r}")
    rows = []
    for values in rows_iter:
        if all(v is None for v in values):
            continue
        rows.append(dict(zip(header, values)))
    return rows


def parse_txt_bytes(content: bytes) -> str:
    """Decode a standalone .txt/.md file's bytes into plain text.

    Runs `ftfy.fix_text()` on the result -- already a project dependency
    for exactly this reason (see AGENTS.md's V7 pilot mojibake note): a
    document dropped in by a researcher may already be
    Windows-1252-decoded-as-something-else corrupted before it ever
    reaches Decifra, and fixing that once at import time is cheaper than
    every downstream consumer having to guard against it.
    """
    return ftfy.fix_text(content.decode("utf-8-sig"))


def parse_docx_bytes(content: bytes) -> str:
    """Extract plain text from a .docx file: every paragraph, then every
    table cell (tables don't appear in `Document.paragraphs` at all), each
    joined by a blank line. Formatting (bold, headings, etc.) is discarded
    -- Decifra codes text content, not document structure."""
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n\n".join(parts)


def parse_pdf_bytes(content: bytes) -> str:
    """Extract plain text from a PDF, page by page, joined by a blank
    line. No OCR: a scanned/image-only PDF yields empty or near-empty
    text per page, silently -- out of scope per AGENTS.md's explicitly
    deferred "image/scanned-PDF extraction (OCR)"."""
    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())
