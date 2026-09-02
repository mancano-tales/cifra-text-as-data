import io

import docx
import openpyxl
import pytest

from text_as_data.corpus_import import (
    parse_csv_rows,
    parse_docx_bytes,
    parse_pdf_bytes,
    parse_txt_bytes,
    parse_xlsx_rows,
)


def test_parse_csv_rows_reads_header_and_rows():
    content = "title,body\nA,First doc\nB,Second doc\n".encode("utf-8")

    rows = parse_csv_rows(content)

    assert rows == [
        {"title": "A", "body": "First doc"},
        {"title": "B", "body": "Second doc"},
    ]


def test_parse_csv_rows_strips_utf8_bom_from_the_first_header():
    content = "﻿title,body\nA,First doc\n".encode("utf-8")

    rows = parse_csv_rows(content)

    assert rows == [{"title": "A", "body": "First doc"}]


def _make_xlsx_bytes(rows: list[list]) -> bytes:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def test_parse_xlsx_rows_reads_header_and_rows():
    content = _make_xlsx_bytes([["title", "body"], ["A", "First doc"], ["B", "Second doc"]])

    rows = parse_xlsx_rows(content)

    assert rows == [
        {"title": "A", "body": "First doc"},
        {"title": "B", "body": "Second doc"},
    ]


def test_parse_xlsx_rows_skips_blank_trailing_rows():
    content = _make_xlsx_bytes([["title", "body"], ["A", "First doc"], [None, None]])

    rows = parse_xlsx_rows(content)

    assert len(rows) == 1


def test_parse_xlsx_rows_keeps_a_row_with_a_blank_first_cell_when_other_cells_have_data():
    # A blank-row guard keyed only on the first column silently drops real
    # data rows whenever that column holds optional metadata (e.g. "notes")
    # left blank on some rows while the actual text lives in column B.
    content = _make_xlsx_bytes([["notes", "body"], [None, "First doc"], ["some note", "Second doc"]])

    rows = parse_xlsx_rows(content)

    assert rows == [
        {"notes": None, "body": "First doc"},
        {"notes": "some note", "body": "Second doc"},
    ]


def test_parse_xlsx_rows_rejects_a_completely_empty_workbook():
    content = _make_xlsx_bytes([])

    with pytest.raises(ValueError, match="no header row"):
        parse_xlsx_rows(content)


def test_parse_xlsx_rows_rejects_an_unnamed_header_column():
    content = _make_xlsx_bytes([["title", None], ["A", "First doc"]])

    with pytest.raises(ValueError, match="unnamed cells"):
        parse_xlsx_rows(content)


def test_parse_txt_bytes_decodes_utf8_with_bom():
    content = "﻿Hello, world.".encode("utf-8")

    text = parse_txt_bytes(content)

    assert text == "Hello, world."


def test_parse_txt_bytes_fixes_mojibake():
    # "instituições" mis-decoded as Windows-1252 read as UTF-8 then
    # re-encoded, the exact corruption pattern AGENTS.md's V7 pilot notes
    # describe -- ftfy.fix_text() should undo it.
    broken = "instituiÃ§Ãµes".encode("utf-8")

    text = parse_txt_bytes(broken)

    assert text == "instituições"


def _make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values):
                cell.text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_parse_docx_bytes_extracts_paragraph_text():
    content = _make_docx_bytes(["First paragraph.", "Second paragraph."])

    text = parse_docx_bytes(content)

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_parse_docx_bytes_extracts_table_cell_text():
    content = _make_docx_bytes(["Intro."], table_rows=[["Header A", "Header B"], ["Value 1", "Value 2"]])

    text = parse_docx_bytes(content)

    assert "Header A" in text
    assert "Value 2" in text


def _make_pdf_bytes(text: str) -> bytes:
    """A hand-built, minimal-but-valid single-page PDF with one text run --
    pypdf can read this without a real PDF-writing library as a
    dependency. Includes a correct xref table (not just %PDF.../%%EOF),
    which pypdf requires to parse at all."""
    objects = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/MediaBox[0 0 200 200]/Contents 5 0 R>>",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    stream = f"BT /F1 12 Tf 10 100 Td ({text}) Tj ET".encode()
    objects.append(b"<</Length " + str(len(stream)).encode() + b">>\nstream\n" + stream + b"\nendstream")

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n"
    xref_offset = len(out)
    out += b"xref\n0 " + str(len(objects) + 1).encode() + b"\n0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += ("%010d 00000 n \n" % offset).encode()
    out += b"trailer\n<</Size " + str(len(objects) + 1).encode() + b"/Root 1 0 R>>\n"
    out += b"startxref\n" + str(xref_offset).encode() + b"\n%%EOF"
    return bytes(out)


def test_parse_pdf_bytes_extracts_text():
    content = _make_pdf_bytes("Hello from a PDF")

    text = parse_pdf_bytes(content)

    assert "Hello from a PDF" in text
